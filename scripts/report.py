"""
Word (.docx) report generator.

`generate_report` produces a self-contained `.docx` for one audit run. If
`templates/report_template.docx` exists it's used as a starting point;
otherwise a clean default layout is built programmatically.

Charts are pre-rendered as PNGs (see `lib/charts.py`) and inserted at
fixed locations.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Mapping

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from lib.charts import render_all
from lib.normalize import AuditInput, DimensionResult, Finding


SEVERITY_ICONS = {
    "critical": "🟥",
    "warning": "🟧",
    "info": "🟦",
    "positive": "🟩",
}


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 0:
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xBA)


def _add_score_badge(doc: Document, overall: float, grade: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{overall:.1f} / 100")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xBA)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Grade: {grade}")
    r2.font.size = Pt(18)
    r2.font.bold = True


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in (table.cell(i, 0), table.cell(i, 1)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.autofit = True


def _add_score_breakdown_table(
    doc: Document,
    results: Mapping[str, DimensionResult],
    weights: Mapping[str, float],
    overall: float,
) -> None:
    table = doc.add_table(rows=len(results) + 2, cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["Dimension", "Score", "Weight", "Weighted", "Key finding"]
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    for i, (name, result) in enumerate(results.items(), start=1):
        weight = weights.get(name, 0.0)
        weighted = result.score * weight
        first_finding = next(
            (f for f in result.findings if f.severity != "positive"),
            result.findings[0] if result.findings else None,
        )
        key_finding = first_finding.title if first_finding else "—"
        table.cell(i, 0).text = name.title()
        table.cell(i, 1).text = f"{result.score:.0f}"
        table.cell(i, 2).text = f"{weight * 100:.0f}%"
        table.cell(i, 3).text = f"{weighted:.1f}"
        table.cell(i, 4).text = key_finding

    total_row = len(results) + 1
    table.cell(total_row, 0).text = "TOTAL"
    table.cell(total_row, 3).text = f"{overall:.1f}"
    for c in range(5):
        for run in table.cell(total_row, c).paragraphs[0].runs:
            run.font.bold = True


def _prioritized_findings(results: Mapping[str, DimensionResult]) -> list[tuple[str, Finding]]:
    """Flatten findings across dimensions, sorted by priority (impact × ease) desc."""
    flat: list[tuple[str, Finding]] = []
    for name, result in results.items():
        for f in result.findings:
            if f.severity == "positive":
                continue
            flat.append((name, f))
    flat.sort(key=lambda pair: pair[1].priority_score, reverse=True)
    return flat


def _add_findings_list(
    doc: Document,
    findings: list[tuple[str, Finding]],
    max_items: int = 8,
) -> None:
    for dim, f in findings[:max_items]:
        p = doc.add_paragraph()
        run = p.add_run(f"{SEVERITY_ICONS.get(f.severity, '·')} {f.title}")
        run.font.bold = True
        run.font.size = Pt(11)
        evidence = doc.add_paragraph(f"Evidence: {f.evidence}")
        evidence.paragraph_format.left_indent = Inches(0.3)
        if f.recommended_action:
            action = doc.add_paragraph(f"Action: {f.recommended_action}")
            action.paragraph_format.left_indent = Inches(0.3)
        meta = doc.add_paragraph(f"[{dim}] · Impact: {f.impact} · Ease: {f.ease}")
        meta.paragraph_format.left_indent = Inches(0.3)
        for run in meta.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _add_dimension_section(doc: Document, result: DimensionResult, charts: Mapping[str, Path]) -> None:
    _add_heading(doc, f"{result.name.title()} — {result.score:.0f}/100", level=2)

    # Metrics table
    if result.metrics:
        rows = []
        for k, v in result.metrics.items():
            if v is None:
                rows.append((k.replace("_", " ").title(), "—"))
            elif isinstance(v, float):
                rows.append((k.replace("_", " ").title(), f"{v:.2f}"))
            else:
                rows.append((k.replace("_", " ").title(), str(v)))
        _add_kv_table(doc, rows)

    # Findings
    if result.findings:
        doc.add_paragraph().add_run("Findings:").font.bold = True
        for f in result.findings:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{SEVERITY_ICONS.get(f.severity, '·')} {f.title}")
            run.font.bold = True
            doc.add_paragraph(f"  {f.evidence}").paragraph_format.left_indent = Inches(0.3)
            if f.recommended_action:
                doc.add_paragraph(f"  → {f.recommended_action}").paragraph_format.left_indent = Inches(0.3)

    # Embed chart for relevant dimensions
    chart_map = {
        "engagement": "engagement_over_time",
        "cadence": "posting_heatmap",
        "hashtags": "hashtag_top",
    }
    chart_key = chart_map.get(result.name)
    if chart_key and chart_key in charts:
        try:
            doc.add_picture(str(charts[chart_key]), width=Inches(6.0))
        except Exception:
            pass


def generate_report(
    *,
    account: str,
    audit_input: AuditInput,
    results: Mapping[str, DimensionResult],
    overall_score: float,
    grade: str,
    weights: Mapping[str, float],
    output_dir: str | Path = "./output",
    template_path: str | Path | None = None,
) -> Path:
    """Generate the Word report. Returns the absolute path to the saved .docx."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Render charts to a per-run subdirectory
    today = date.today().isoformat()
    chart_dir = output_dir / f"charts_{account}_{today}"
    dim_scores = {name: r.score for name, r in results.items()}
    charts = render_all(audit_input, dim_scores, chart_dir)

    # Use the hand-built template if it's there, else start fresh
    template = Path(template_path) if template_path else None
    if template and template.exists():
        doc = Document(str(template))
        # Hand-built template support — placeholder replacement (Phase 1: minimal)
        # For now we still append the standard sections after the template content;
        # full placeholder substitution is a follow-up improvement once you build the .docx.
        doc.add_page_break()
    else:
        doc = Document()

    # ---- Cover ----
    _add_heading(doc, f"Instagram Audit — {audit_input.profile.display_name}", level=0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"@{audit_input.profile.username}  ·  "
        f"{audit_input.period_start.isoformat()} → {audit_input.period_end.isoformat()}"
    ).font.size = Pt(11)
    _add_score_badge(doc, overall_score, grade)
    doc.add_paragraph(
        f"Generated: {today} · Source: {audit_input.source.upper()} · "
        f"Followers: {audit_input.profile.follower_count:,} · "
        f"Posts in period: {len(audit_input.posts)}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ---- Executive summary ----
    _add_heading(doc, "Executive Summary", level=1)
    summary = _build_executive_summary(audit_input, results, overall_score, grade)
    for para in summary:
        doc.add_paragraph(para)

    # ---- Score breakdown ----
    _add_heading(doc, "Score Breakdown", level=1)
    _add_score_breakdown_table(doc, results, weights, overall_score)
    if "score_radar" in charts:
        doc.add_paragraph()
        try:
            doc.add_picture(str(charts["score_radar"]), width=Inches(5.5))
        except Exception:
            pass

    # ---- Top findings ----
    _add_heading(doc, "Top Findings", level=1)
    prioritized = _prioritized_findings(results)
    _add_findings_list(doc, prioritized, max_items=8)

    # ---- Per-dimension detail ----
    doc.add_page_break()
    _add_heading(doc, "Detail by Dimension", level=1)
    for name, result in results.items():
        _add_dimension_section(doc, result, charts)

    # ---- Action plan ----
    doc.add_page_break()
    _add_heading(doc, "Prioritized Action Plan", level=1)
    if prioritized:
        action_table = doc.add_table(rows=min(11, len(prioritized) + 1), cols=5)
        action_table.style = "Light Grid Accent 1"
        for c, h in enumerate(["#", "Action", "Impact", "Ease", "Dimension"]):
            cell = action_table.cell(0, c)
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        for i, (dim, f) in enumerate(prioritized[:10], start=1):
            action_table.cell(i, 0).text = str(i)
            action_table.cell(i, 1).text = f.recommended_action or f.title
            action_table.cell(i, 2).text = f.impact.title()
            action_table.cell(i, 3).text = f.ease.title()
            action_table.cell(i, 4).text = dim.title()

    # ---- Methodology ----
    _add_heading(doc, "Methodology", level=1)
    doc.add_paragraph(
        f"Audit period: {audit_input.period_start.isoformat()} to "
        f"{audit_input.period_end.isoformat()}. "
        f"Source: {audit_input.source.upper()}. "
        f"Scoring weights: see references/scoring_weights.json. "
        f"Tune weights and thresholds without code changes by editing that file."
    )
    doc.add_paragraph(
        "Generated by Instagram Audit Skill — "
        "github.com/TapasDas1982/Instagram_Audit_Skill"
    ).runs[0].font.size = Pt(9)

    # Save
    out_path = output_dir / f"{account}_{today}.docx"
    doc.save(str(out_path))
    return out_path.resolve()


def _build_executive_summary(
    audit_input: AuditInput,
    results: Mapping[str, DimensionResult],
    overall_score: float,
    grade: str,
) -> list[str]:
    """Three short paragraphs: headline, what's working, what's hurting."""
    # Headline
    descriptor = {
        "A": "excellent",
        "B": "solid",
        "C": "average",
        "D": "below average",
        "F": "critical",
    }.get(grade, "")
    headline = (
        f"@{audit_input.profile.username} scored {overall_score:.0f}/100 "
        f"({grade} — {descriptor}) over the {(audit_input.period_end - audit_input.period_start).days + 1}-day "
        f"window. {len(audit_input.posts)} posts published; {audit_input.profile.follower_count:,} followers."
    )

    # Strengths
    strengths = sorted(results.items(), key=lambda kv: kv[1].score, reverse=True)[:2]
    strength_text = "Strongest dimensions: " + ", ".join(
        f"{name.title()} ({r.score:.0f}/100)" for name, r in strengths
    ) + "."

    # Weaknesses
    weaknesses = sorted(results.items(), key=lambda kv: kv[1].score)[:2]
    weakness_text = "Biggest gaps: " + ", ".join(
        f"{name.title()} ({r.score:.0f}/100)" for name, r in weaknesses
    ) + "."

    # Top 3 actions preview
    prioritized = _prioritized_findings(results)
    if prioritized:
        actions = "Top 3 actions: " + " | ".join(
            f"{i + 1}. {f.recommended_action or f.title}"
            for i, (_, f) in enumerate(prioritized[:3])
            if f.recommended_action
        )
    else:
        actions = ""

    return [headline, strength_text, weakness_text] + ([actions] if actions else [])
