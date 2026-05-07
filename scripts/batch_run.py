"""
Monthly batch audit runner.

Iterates all active owned accounts in the `accounts` table, runs a full audit
for each, and produces:
  - output/batch_YYYY-MM-DD/
      ├── {username}_{date}.docx      (per-account detail report)
      └── batch_summary_{date}.docx  (one-page cover summary)

Usage:
    python scripts/batch_run.py [--output-dir ./output] [--period-days 30] [--no-email]
    python scripts/batch_run.py --accounts twistnturns,twistnturns_salt_lake
"""

from __future__ import annotations

import argparse
import logging
import subprocess
import sys
from contextlib import contextmanager
from datetime import date
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# Project root — scripts/ lives one level below the project root
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent
SCRIPTS_DIR = PROJECT_ROOT / "scripts"

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Config helpers
# ---------------------------------------------------------------------------

def _load_config() -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    """Return (MYSQL, BREVO, AUDIT) from config/config.py.

    If the file doesn't exist, raises ImportError with a helpful message.
    """
    sys.path.insert(0, str(PROJECT_ROOT / "config"))
    try:
        import config as cfg  # type: ignore[import]
    except ImportError as exc:
        raise ImportError(
            "config/config.py not found. "
            "Copy config/config.example.py to config/config.py and fill in your values."
        ) from exc
    return cfg.MYSQL, cfg.BREVO, cfg.AUDIT


# ---------------------------------------------------------------------------
# Database helpers
# ---------------------------------------------------------------------------

@contextmanager
def _db_connect(mysql_config: Mapping[str, Any]):
    """Context manager — yields a raw mysql.connector connection."""
    try:
        import mysql.connector  # type: ignore[import]
    except ImportError as exc:
        raise RuntimeError(
            "mysql-connector-python is not installed. "
            "Run: pip install mysql-connector-python"
        ) from exc

    conn = mysql.connector.connect(
        host=mysql_config["host"],
        port=mysql_config.get("port", 3306),
        user=mysql_config["user"],
        password=mysql_config["password"],
        database=mysql_config["database"],
    )
    try:
        yield conn
    finally:
        conn.close()


def _fetch_active_accounts(mysql_config: Mapping[str, Any]) -> list[dict[str, Any]]:
    """Return all active owned accounts as a list of dicts."""
    with _db_connect(mysql_config) as conn:
        cur = conn.cursor(dictionary=True)
        cur.execute(
            "SELECT id, username, studio_location "
            "FROM accounts "
            "WHERE account_type = 'owned' AND is_active = 1 "
            "ORDER BY studio_location, username"
        )
        return list(cur.fetchall())


def _get_previous_score(
    mysql_config: Mapping[str, Any],
    account_id: int,
) -> float | None:
    """Return the second-most-recent overall_score for an account, or None."""
    try:
        with _db_connect(mysql_config) as conn:
            cur = conn.cursor()
            cur.execute(
                "SELECT overall_score FROM audits "
                "WHERE account_id = %s "
                "ORDER BY audit_date DESC LIMIT 2",
                (account_id,),
            )
            rows = cur.fetchall()
            if len(rows) < 2:
                return None
            val = rows[1][0]
            return float(val) if val is not None else None
    except Exception as exc:  # non-fatal
        logger.warning("Could not fetch previous score for account %s: %s", account_id, exc)
        return None


def _get_current_score_from_db(
    mysql_config: Mapping[str, Any],
    account_id: int,
) -> float | None:
    """Return the most-recent overall_score for an account from the DB, or None."""
    try:
        with _db_connect(mysql_config) as conn:
            cur = conn.cursor()
            cur.execute(
                "SELECT overall_score FROM audits "
                "WHERE account_id = %s "
                "ORDER BY audit_date DESC LIMIT 1",
                (account_id,),
            )
            row = cur.fetchone()
            if row is None:
                return None
            return float(row[0]) if row[0] is not None else None
    except Exception as exc:  # non-fatal
        logger.warning("Could not fetch current score for account %s: %s", account_id, exc)
        return None


# ---------------------------------------------------------------------------
# Audit subprocess runner
# ---------------------------------------------------------------------------

def _run_single_audit(
    username: str,
    location: str | None,
    batch_dir: Path,
    period_days: int,
) -> dict[str, Any]:
    """Run audit.py for one account. Returns a result dict.

    Never raises — all errors are captured and returned in result["error"].
    """
    cmd = [
        sys.executable,
        str(SCRIPTS_DIR / "audit.py"),
        "--source", "api",
        "--account", username,
        "--output-dir", str(batch_dir),
        "--period-days", str(period_days),
    ]
    if location:
        cmd += ["--studio-location", location]

    logger.info("Running audit for @%s …", username)
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    except subprocess.TimeoutExpired:
        return {"success": False, "report_path": None, "score": None,
                "error": "Audit timed out after 300 seconds."}
    except Exception as exc:
        return {"success": False, "report_path": None, "score": None,
                "error": f"Subprocess exception: {exc}"}

    if proc.returncode != 0:
        err_snippet = (proc.stderr or "").strip()[-500:]  # last 500 chars
        logger.warning("Audit failed for @%s (exit %d): %s", username, proc.returncode, err_snippet)
        return {"success": False, "report_path": None, "score": None,
                "error": err_snippet or f"Exit code {proc.returncode}"}

    # Find the report file — audit.py writes {username}_{today}.docx
    today_iso = date.today().isoformat()
    expected_report = batch_dir / f"{username}_{today_iso}.docx"
    report_path = str(expected_report) if expected_report.exists() else None

    logger.info("Audit complete for @%s — report: %s", username, report_path)
    return {"success": True, "report_path": report_path, "score": None, "error": None}


# ---------------------------------------------------------------------------
# Trend helper
# ---------------------------------------------------------------------------

def _trend_arrow(current: float | None, previous: float | None) -> str:
    """Return a Unicode trend arrow comparing current vs previous score."""
    if current is None or previous is None:
        return "—"  # em-dash: no data
    diff = current - previous
    if diff > 2.0:
        return "↑"  # ↑
    if diff < -2.0:
        return "↓"  # ↓
    return "→"  # →


# ---------------------------------------------------------------------------
# Summary .docx builder
# ---------------------------------------------------------------------------

def _build_summary_docx(
    results: dict[str, dict[str, Any]],
    batch_dir: Path,
    month_label: str,
    today_iso: str,
) -> Path:
    """Write batch_summary_{date}.docx and return its Path."""
    doc = Document()

    # Title
    title = doc.add_heading("Twist N Turns — Monthly Instagram Audit Summary", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xBA)

    # Subtitle
    total = len(results)
    passed = sum(1 for r in results.values() if r["success"])
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"Generated: {today_iso}  ·  Audit month: {month_label}  ·  "
        f"Accounts audited: {passed}/{total}"
    ).font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Summary table
    col_headers = ["Location", "Account", "Score", "vs Last Month", "Status", "Report File"]
    n_data_rows = max(len(results), 1)
    table = doc.add_table(rows=n_data_rows + 1, cols=len(col_headers))
    table.style = "Light Grid Accent 1"

    # Header row
    for c, h in enumerate(col_headers):
        cell = table.cell(0, c)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    # Data rows
    for i, (username, info) in enumerate(results.items(), start=1):
        score = info.get("score")
        score_str = f"{score:.1f}" if score is not None else "—"
        trend = info.get("trend_arrow", "—")
        status = "✓" if info["success"] else "✗"
        report_file = Path(info["report_path"]).name if info.get("report_path") else "—"
        location = info.get("studio_location") or "—"

        table.cell(i, 0).text = location
        table.cell(i, 1).text = f"@{username}"
        table.cell(i, 2).text = score_str
        table.cell(i, 3).text = trend
        table.cell(i, 4).text = status
        table.cell(i, 5).text = report_file

    doc.add_paragraph()  # spacer
    footer = doc.add_paragraph("Full per-account reports are in the batch directory.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    out_path = batch_dir / f"batch_summary_{today_iso}.docx"
    doc.save(str(out_path))
    logger.info("Batch summary written to %s", out_path)
    return out_path


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Run monthly Instagram audits for all active owned accounts."
    )
    parser.add_argument(
        "--output-dir",
        default="./output",
        help="Root output directory. A batch_YYYY-MM-DD/ subdirectory is created automatically.",
    )
    parser.add_argument(
        "--period-days",
        type=int,
        default=None,
        help="Number of days in the audit window (overrides AUDIT config).",
    )
    parser.add_argument(
        "--no-email",
        action="store_true",
        help="Skip sending the batch summary email.",
    )
    parser.add_argument(
        "--accounts",
        default=None,
        help="Comma-separated list of usernames to audit (default: all active owned accounts).",
    )
    return parser.parse_args()


def main() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
    )

    args = _parse_args()
    mysql_config, brevo_config, audit_config = _load_config()

    period_days: int = args.period_days or audit_config.get("default_period_days", 30)
    today_iso = date.today().isoformat()
    month_label = today_iso[:7]  # "YYYY-MM"

    batch_dir = Path(args.output_dir) / f"batch_{today_iso}"
    batch_dir.mkdir(parents=True, exist_ok=True)
    logger.info("Batch output directory: %s", batch_dir.resolve())

    # Determine which accounts to run
    if args.accounts:
        requested = {u.strip() for u in args.accounts.split(",")}
    else:
        requested = None

    all_accounts = _fetch_active_accounts(mysql_config)
    if requested is not None:
        accounts = [a for a in all_accounts if a["username"] in requested]
        missing = requested - {a["username"] for a in accounts}
        if missing:
            logger.warning(
                "These usernames were not found in the DB (or are not active owned accounts): %s",
                ", ".join(sorted(missing)),
            )
    else:
        accounts = all_accounts

    if not accounts:
        logger.error("No active owned accounts found — nothing to audit.")
        sys.exit(1)

    logger.info("Auditing %d account(s).", len(accounts))

    # Run audits
    results: dict[str, dict[str, Any]] = {}
    for account in accounts:
        username: str = account["username"]
        location: str | None = account.get("studio_location")
        account_id: int = account["id"]

        result = _run_single_audit(username, location, batch_dir, period_days)
        result["studio_location"] = location
        result["account_id"] = account_id

        # Fetch current score from DB after the audit wrote it
        current_score = _get_current_score_from_db(mysql_config, account_id)
        prev_score = _get_previous_score(mysql_config, account_id)
        result["score"] = current_score
        result["trend_arrow"] = _trend_arrow(current_score, prev_score)

        results[username] = result

    # Build summary .docx
    summary_path = _build_summary_docx(results, batch_dir, month_label, today_iso)

    # Email
    if not args.no_email:
        from lib.mailer import send_batch_email  # local import avoids import-time side effects
        ok = send_batch_email(
            brevo_config=brevo_config,
            batch_dir=batch_dir,
            results=results,
            month_label=month_label,
        )
        if ok:
            logger.info("Batch email sent successfully.")
        else:
            logger.warning("Batch email was NOT sent (see warnings above).")

    # Print human-readable summary to stdout
    passed = [u for u, r in results.items() if r["success"]]
    failed = [u for u, r in results.items() if not r["success"]]

    print(f"\n{'=' * 60}")
    print(f"  Batch complete — {today_iso}  |  Month: {month_label}")
    print(f"  Audited: {len(results)}  |  OK: {len(passed)}  |  Failed: {len(failed)}")
    print(f"  Summary report: {summary_path}")
    print(f"{'=' * 60}")

    for username, info in results.items():
        score_str = f"{info['score']:.1f}" if info.get("score") is not None else "n/a"
        trend = info.get("trend_arrow", "—")
        status_str = "OK" if info["success"] else f"FAIL: {info.get('error', '')[:80]}"
        print(f"  {'OK  ' if info['success'] else 'FAIL'}  @{username:<30} score={score_str:<7} {trend}  {status_str}")

    if failed:
        print(f"\n  {len(failed)} account(s) failed. Check logs for details.")
        sys.exit(2)


if __name__ == "__main__":
    main()
